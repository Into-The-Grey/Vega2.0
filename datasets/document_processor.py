"""
Document processing utilities for Vega 2.0 multi-modal learning.

This module provides:
- Document format detection and validation
- Text extraction from multiple document formats (PDF, DOCX, TXT, RTF, HTML)
- Document sanitization and security checks
- Metadata extraction and preservation
"""

import os
import mimetypes
from typing import Dict, List, Optional, Tuple, Any, Union
from pathlib import Path
import logging
from dataclasses import dataclass, asdict
from datetime import datetime
import tempfile
import shutil

# Document processing imports
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. Install with: pip install PyPDF2")

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with: pip install python-docx")

try:
    from bs4 import BeautifulSoup

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logging.warning(
        "BeautifulSoup4 not available. Install with: pip install beautifulsoup4"
    )

try:
    from striprtf.striprtf import rtf_to_text

    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False
    logging.warning("striprtf not available. Install with: pip install striprtf")

try:
    import magic

    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logging.warning(
        "python-magic not available. Install with: pip install python-magic"
    )

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {
    ".pdf": "PDF",
    ".docx": "Word Document",
    ".doc": "Word Document (Legacy)",
    ".txt": "Text",
    ".rtf": "Rich Text Format",
    ".html": "HTML",
    ".htm": "HTML",
    ".md": "Markdown",
}


@dataclass
class DocumentMetadata:
    """Document metadata container."""

    file_path: str
    file_size: int
    format_type: str
    mime_type: str
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    language: Optional[str] = None
    encoding: Optional[str] = None


@dataclass
class DocumentContent:
    """Container for extracted document content."""

    text: str
    metadata: DocumentMetadata
    formatting: Optional[Dict[str, Any]] = None
    structure: Optional[Dict[str, Any]] = None
    extracted_data: Optional[Dict[str, Any]] = None
    processing_errors: Optional[List[str]] = None


class DocumentValidator:
    """Validate and sanitize document files."""

    MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def __del__(self):
        if hasattr(self, "temp_dir") and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def validate_document(self, file_path: str) -> Tuple[bool, List[str]]:
        """Validate document file for security and format compliance."""
        errors = []

        try:
            # Check file existence
            if not os.path.exists(file_path):
                errors.append("File does not exist")
                return False, errors

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.MAX_FILE_SIZE:
                errors.append(
                    f"File size ({file_size} bytes) exceeds maximum ({self.MAX_FILE_SIZE} bytes)"
                )
                return False, errors

            if file_size == 0:
                errors.append("File is empty")
                return False, errors

            # Check file extension
            ext = Path(file_path).suffix.lower()
            if ext not in SUPPORTED_FORMATS:
                errors.append(f"Unsupported file format: {ext}")
                return False, errors

            # MIME type validation
            mime_type = self._get_mime_type(file_path)
            if not self._validate_mime_type(ext, mime_type):
                errors.append(
                    f"MIME type {mime_type} doesn't match file extension {ext}"
                )

            # Basic content validation
            if not self._validate_file_content(file_path, ext):
                errors.append("File content validation failed")
                return False, errors

            return len(errors) == 0, errors

        except Exception as e:
            errors.append(f"Validation error: {e}")
            return False, errors

    def _get_mime_type(self, file_path: str) -> str:
        """Get MIME type of file."""
        if MAGIC_AVAILABLE:
            try:
                return magic.from_file(file_path, mime=True)
            except Exception:
                pass

        # Fallback to mimetypes
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or "application/octet-stream"

    def _validate_mime_type(self, extension: str, mime_type: str) -> bool:
        """Validate MIME type matches file extension."""
        expected_mimes = {
            ".pdf": ["application/pdf"],
            ".docx": [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ],
            ".doc": ["application/msword"],
            ".txt": ["text/plain"],
            ".rtf": ["application/rtf", "text/rtf"],
            ".html": ["text/html"],
            ".htm": ["text/html"],
            ".md": ["text/markdown", "text/x-markdown", "text/plain"],
        }

        expected = expected_mimes.get(extension, [])
        return mime_type in expected if expected else True

    def _validate_file_content(self, file_path: str, extension: str) -> bool:
        """Basic file content validation."""
        try:
            if extension == ".pdf":
                return self._validate_pdf_content(file_path)
            elif extension == ".docx":
                return self._validate_docx_content(file_path)
            elif extension in [".txt", ".md"]:
                return self._validate_text_content(file_path)
            elif extension in [".html", ".htm"]:
                return self._validate_html_content(file_path)
            elif extension == ".rtf":
                return self._validate_rtf_content(file_path)
            return True
        except Exception as e:
            logger.warning(f"Content validation failed: {e}")
            return False

    def _validate_pdf_content(self, file_path: str) -> bool:
        """Validate PDF file content."""
        if not PDF_AVAILABLE:
            return True

        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                return len(reader.pages) > 0
        except Exception:
            return False

    def _validate_docx_content(self, file_path: str) -> bool:
        """Validate DOCX file content."""
        if not DOCX_AVAILABLE:
            return True

        try:
            doc = DocxDocument(file_path)
            return len(doc.paragraphs) >= 0  # Can be empty but should be readable
        except Exception:
            return False

    def _validate_text_content(self, file_path: str) -> bool:
        """Validate text file content."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(1024)  # Read first 1KB
                return isinstance(content, str)
        except Exception:
            return False

    def _validate_html_content(self, file_path: str) -> bool:
        """Validate HTML file content."""
        if not HTML_AVAILABLE:
            return True

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(1024)
                soup = BeautifulSoup(content, "html.parser")
                return soup is not None
        except Exception:
            return False

    def _validate_rtf_content(self, file_path: str) -> bool:
        """Validate RTF file content."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read(100)
                return content.startswith("{\\rtf")
        except Exception:
            return False


class DocumentProcessor:
    """Main document processing engine."""

    def __init__(self):
        self.validator = DocumentValidator()
        self.temp_dir = tempfile.mkdtemp()

    def __del__(self):
        if hasattr(self, "temp_dir") and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def process_document(
        self,
        file_path: str,
        extract_structure: bool = True,
        extract_metadata: bool = True,
    ) -> DocumentContent:
        """Process a document and extract all relevant information."""

        # Validate document
        is_valid, errors = self.validator.validate_document(file_path)
        if not is_valid:
            return DocumentContent(
                text="",
                metadata=DocumentMetadata(
                    file_path=file_path,
                    file_size=0,
                    format_type="unknown",
                    mime_type="unknown",
                ),
                processing_errors=errors,
            )

        # Extract content based on format
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                return self._process_pdf(file_path, extract_structure, extract_metadata)
            elif ext == ".docx":
                return self._process_docx(
                    file_path, extract_structure, extract_metadata
                )
            elif ext in [".txt", ".md"]:
                return self._process_text(
                    file_path, extract_structure, extract_metadata
                )
            elif ext in [".html", ".htm"]:
                return self._process_html(
                    file_path, extract_structure, extract_metadata
                )
            elif ext == ".rtf":
                return self._process_rtf(file_path, extract_structure, extract_metadata)
            else:
                return DocumentContent(
                    text="",
                    metadata=DocumentMetadata(
                        file_path=file_path,
                        file_size=os.path.getsize(file_path),
                        format_type=SUPPORTED_FORMATS.get(ext, "unknown"),
                        mime_type=self.validator._get_mime_type(file_path),
                    ),
                    processing_errors=[f"Unsupported format: {ext}"],
                )

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return DocumentContent(
                text="",
                metadata=DocumentMetadata(
                    file_path=file_path,
                    file_size=(
                        os.path.getsize(file_path) if os.path.exists(file_path) else 0
                    ),
                    format_type=SUPPORTED_FORMATS.get(ext, "unknown"),
                    mime_type=self.validator._get_mime_type(file_path),
                ),
                processing_errors=[f"Processing error: {e}"],
            )

    def _create_base_metadata(
        self, file_path: str, format_type: str
    ) -> DocumentMetadata:
        """Create base metadata for a document."""
        stat = os.stat(file_path)

        return DocumentMetadata(
            file_path=file_path,
            file_size=stat.st_size,
            format_type=format_type,
            mime_type=self.validator._get_mime_type(file_path),
            created_date=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            modified_date=datetime.fromtimestamp(stat.st_mtime).isoformat(),
        )

    def _process_pdf(
        self, file_path: str, extract_structure: bool, extract_metadata: bool
    ) -> DocumentContent:
        """Process PDF document."""
        if not PDF_AVAILABLE:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "PDF"),
                processing_errors=["PyPDF2 not available"],
            )

        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)

                # Extract text
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"

                # Create metadata
                metadata = self._create_base_metadata(file_path, "PDF")
                metadata.page_count = len(reader.pages)
                metadata.word_count = len(text.split())
                metadata.character_count = len(text)

                # Extract PDF metadata if available
                if extract_metadata and reader.metadata:
                    metadata.title = reader.metadata.get("/Title")
                    metadata.author = reader.metadata.get("/Author")

                return DocumentContent(
                    text=text.strip(),
                    metadata=metadata,
                    structure=(
                        {"page_count": len(reader.pages)} if extract_structure else None
                    ),
                )

        except Exception as e:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "PDF"),
                processing_errors=[f"PDF processing error: {e}"],
            )

    def _process_docx(
        self, file_path: str, extract_structure: bool, extract_metadata: bool
    ) -> DocumentContent:
        """Process DOCX document."""
        if not DOCX_AVAILABLE:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "Word Document"),
                processing_errors=["python-docx not available"],
            )

        try:
            doc = DocxDocument(file_path)

            # Extract text
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Create metadata
            metadata = self._create_base_metadata(file_path, "Word Document")
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)

            # Extract document properties if available
            if extract_metadata and hasattr(doc, "core_properties"):
                props = doc.core_properties
                metadata.title = props.title
                metadata.author = props.author
                if props.created:
                    metadata.created_date = props.created.isoformat()
                if props.modified:
                    metadata.modified_date = props.modified.isoformat()

            # Extract structure
            structure = None
            if extract_structure:
                structure = {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                }

            return DocumentContent(
                text=text.strip(), metadata=metadata, structure=structure
            )

        except Exception as e:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "Word Document"),
                processing_errors=[f"DOCX processing error: {e}"],
            )

    def _process_text(
        self, file_path: str, extract_structure: bool, extract_metadata: bool
    ) -> DocumentContent:
        """Process plain text document."""
        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
            text = ""
            encoding_used = "utf-8"

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue

            if not text:
                raise ValueError("Could not decode file with any supported encoding")

            # Create metadata
            metadata = self._create_base_metadata(file_path, "Text")
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)
            metadata.encoding = encoding_used

            # Extract structure
            structure = None
            if extract_structure:
                lines = text.split("\n")
                structure = {
                    "line_count": len(lines),
                    "paragraph_count": len([line for line in lines if line.strip()]),
                }

            return DocumentContent(
                text=text.strip(), metadata=metadata, structure=structure
            )

        except Exception as e:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "Text"),
                processing_errors=[f"Text processing error: {e}"],
            )

    def _process_html(
        self, file_path: str, extract_structure: bool, extract_metadata: bool
    ) -> DocumentContent:
        """Process HTML document."""
        if not HTML_AVAILABLE:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "HTML"),
                processing_errors=["BeautifulSoup4 not available"],
            )

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()

            soup = BeautifulSoup(content, "html.parser")

            # Extract text content
            text = soup.get_text(separator="\n", strip=True)

            # Create metadata
            metadata = self._create_base_metadata(file_path, "HTML")
            metadata.word_count = len(text.split()) if text else 0
            metadata.character_count = len(text) if text else 0

            # Extract HTML metadata
            if extract_metadata:
                title_tag = soup.find("title")
                if title_tag and hasattr(title_tag, "get_text"):
                    metadata.title = title_tag.get_text(strip=True)

                author_meta = soup.find("meta", {"name": "author"})
                if author_meta and hasattr(author_meta, "get"):
                    metadata.author = author_meta.get("content")

            # Extract structure
            structure = None
            if extract_structure:
                try:
                    headings = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
                    paragraphs = soup.find_all("p")
                    links = soup.find_all("a")
                    images = soup.find_all("img")

                    structure = {
                        "heading_count": len(headings) if headings else 0,
                        "paragraph_count": len(paragraphs) if paragraphs else 0,
                        "link_count": len(links) if links else 0,
                        "image_count": len(images) if images else 0,
                    }
                except Exception as struct_e:
                    logger.warning(f"HTML structure extraction failed: {struct_e}")
                    structure = {"extraction_error": str(struct_e)}

            return DocumentContent(
                text=text.strip() if text else "",
                metadata=metadata,
                structure=structure,
                formatting={"html_content": content} if extract_structure else None,
            )

        except Exception as e:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "HTML"),
                processing_errors=[f"HTML processing error: {e}"],
            )

    def _process_rtf(
        self, file_path: str, extract_structure: bool, extract_metadata: bool
    ) -> DocumentContent:
        """Process RTF document."""
        if not RTF_AVAILABLE:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "Rich Text Format"),
                processing_errors=["striprtf not available"],
            )

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                rtf_content = file.read()

            # Convert RTF to plain text
            text = rtf_to_text(rtf_content)

            # Create metadata
            metadata = self._create_base_metadata(file_path, "Rich Text Format")
            metadata.word_count = len(text.split())
            metadata.character_count = len(text)

            return DocumentContent(text=text.strip(), metadata=metadata)

        except Exception as e:
            return DocumentContent(
                text="",
                metadata=self._create_base_metadata(file_path, "Rich Text Format"),
                processing_errors=[f"RTF processing error: {e}"],
            )


# Convenience functions
def process_document(
    file_path: str, extract_structure: bool = True, extract_metadata: bool = True
) -> DocumentContent:
    """Process a single document."""
    processor = DocumentProcessor()
    return processor.process_document(file_path, extract_structure, extract_metadata)


def validate_document(file_path: str) -> Tuple[bool, List[str]]:
    """Validate a document file."""
    validator = DocumentValidator()
    return validator.validate_document(file_path)


def get_supported_formats() -> Dict[str, str]:
    """Get dictionary of supported file formats."""
    return SUPPORTED_FORMATS.copy()


if __name__ == "__main__":
    # Example usage
    test_file = "sample_document.pdf"
    if os.path.exists(test_file):
        result = process_document(test_file)
        print(f"Text: {result.text[:200]}...")
        print(f"Metadata: {result.metadata}")
        if result.processing_errors:
            print(f"Errors: {result.processing_errors}")
