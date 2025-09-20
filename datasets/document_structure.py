"""
Document structure analysis for Vega 2.0 multi-modal learning.

This module provides:
- Hierarchical document structure extraction
- Table detection and extraction
- Header and section analysis
- Text formatting preservation
- Structured data extraction
"""

import re
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass
from pathlib import Path
import logging

# Document processing imports
try:
    import PyPDF2
    import pdfplumber

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup, Tag

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import pandas as pd
    import tabula

    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class TableData:
    """Container for extracted table data."""

    rows: List[List[str]]
    headers: Optional[List[str]] = None
    caption: Optional[str] = None
    position: Optional[Dict[str, Any]] = None
    table_type: Optional[str] = None


@dataclass
class HeaderInfo:
    """Container for header/section information."""

    level: int
    text: str
    position: Optional[int] = None
    page: Optional[int] = None
    subsections: Optional[List["HeaderInfo"]] = None


@dataclass
class DocumentStructure:
    """Container for complete document structure."""

    headers: List[HeaderInfo]
    tables: List[TableData]
    paragraphs: List[Dict[str, Any]]
    lists: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    formatting: Dict[str, Any]
    structure_type: str


class TextAnalyzer:
    """Analyze text structure and patterns."""

    HEADER_PATTERNS = [
        r"^[A-Z][A-Z\s]{3,}$",  # ALL CAPS headers
        r"^\d+\.\s+[A-Z]",  # Numbered sections (1. Introduction)
        r"^[A-Z][a-z]+:",  # Title case with colon
        r"^\w+\.\w+\.",  # Multi-level numbering (1.1.)
        r"^[IVX]+\.\s+[A-Z]",  # Roman numerals
    ]

    LIST_PATTERNS = [
        r"^\s*[-•*]\s+",  # Bullet points
        r"^\s*\d+\.\s+",  # Numbered lists
        r"^\s*[a-zA-Z]\.\s+",  # Lettered lists
        r"^\s*[ivx]+\.\s+",  # Roman numeral lists
    ]

    def __init__(self):
        self.header_regex = [
            re.compile(pattern, re.MULTILINE) for pattern in self.HEADER_PATTERNS
        ]
        self.list_regex = [
            re.compile(pattern, re.MULTILINE) for pattern in self.LIST_PATTERNS
        ]

    def extract_headers(self, text: str) -> List[HeaderInfo]:
        """Extract headers from text using pattern matching."""
        headers = []
        lines = text.split("\n")

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Check for header patterns
            header_level = self._detect_header_level(line)
            if header_level > 0:
                headers.append(HeaderInfo(level=header_level, text=line, position=i))

        return self._build_header_hierarchy(headers)

    def _detect_header_level(self, line: str) -> int:
        """Detect if line is a header and return its level."""
        # Check for numbered sections
        if re.match(r"^\d+\.\s+", line):
            return 1
        elif re.match(r"^\d+\.\d+\.\s+", line):
            return 2
        elif re.match(r"^\d+\.\d+\.\d+\.\s+", line):
            return 3

        # Check for ALL CAPS (likely headers)
        if len(line) > 5 and line.isupper() and not re.search(r"\d", line):
            return 1

        # Check for title case with specific patterns
        if re.match(r"^[A-Z][a-z]+\s+[A-Z]", line) and len(line.split()) <= 6:
            return 2

        # Check for patterns with colons
        if line.endswith(":") and len(line.split()) <= 4:
            return 3

        return 0

    def _build_header_hierarchy(self, headers: List[HeaderInfo]) -> List[HeaderInfo]:
        """Build hierarchical structure from flat header list."""
        if not headers:
            return []

        root_headers = []
        stack = []

        for header in headers:
            # Pop headers from stack that are at same or higher level
            while stack and stack[-1].level >= header.level:
                stack.pop()

            # Add as subsection to parent if stack not empty
            if stack:
                parent = stack[-1]
                if parent.subsections is None:
                    parent.subsections = []
                parent.subsections.append(header)
            else:
                root_headers.append(header)

            stack.append(header)

        return root_headers

    def extract_lists(self, text: str) -> List[Dict[str, Any]]:
        """Extract lists from text."""
        lists = []
        lines = text.split("\n")
        current_list = None

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                if current_list:
                    lists.append(current_list)
                    current_list = None
                continue

            list_match = self._detect_list_item(line)
            if list_match:
                list_type, item_text = list_match

                if current_list is None:
                    current_list = {"type": list_type, "items": [], "start_position": i}
                elif current_list["type"] != list_type:
                    lists.append(current_list)
                    current_list = {"type": list_type, "items": [], "start_position": i}

                current_list["items"].append(item_text)
            else:
                if current_list:
                    lists.append(current_list)
                    current_list = None

        if current_list:
            lists.append(current_list)

        return lists

    def _detect_list_item(self, line: str) -> Optional[Tuple[str, str]]:
        """Detect if line is a list item and return type and content."""
        for pattern in self.list_regex:
            match = pattern.match(line)
            if match:
                marker = match.group()
                text = line[len(marker) :].strip()

                if marker.strip().startswith(("•", "*", "-")):
                    return "bullet", text
                elif re.match(r"\d+\.", marker.strip()):
                    return "numbered", text
                elif re.match(r"[a-zA-Z]\.", marker.strip()):
                    return "lettered", text
                elif re.match(r"[ivx]+\.", marker.strip()):
                    return "roman", text

        return None

    def extract_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """Extract paragraph structure from text."""
        paragraphs = []
        current_paragraph = []
        lines = text.split("\n")

        for i, line in enumerate(lines):
            line_stripped = line.strip()

            if not line_stripped:
                if current_paragraph:
                    paragraphs.append(
                        {
                            "text": " ".join(current_paragraph),
                            "start_line": i - len(current_paragraph),
                            "line_count": len(current_paragraph),
                        }
                    )
                    current_paragraph = []
            else:
                # Skip headers and list items
                if self._detect_header_level(
                    line_stripped
                ) == 0 and not self._detect_list_item(line):
                    current_paragraph.append(line_stripped)

        if current_paragraph:
            paragraphs.append(
                {
                    "text": " ".join(current_paragraph),
                    "start_line": len(lines) - len(current_paragraph),
                    "line_count": len(current_paragraph),
                }
            )

        return paragraphs


class PDFStructureExtractor:
    """Extract structure from PDF documents."""

    def __init__(self):
        self.text_analyzer = TextAnalyzer()

    def extract_structure(self, file_path: str) -> DocumentStructure:
        """Extract complete structure from PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")

        structure_data = {
            "headers": [],
            "tables": [],
            "paragraphs": [],
            "lists": [],
            "metadata": {},
            "formatting": {},
        }

        try:
            # Use pdfplumber for better table extraction if available
            if hasattr(self, "_extract_with_pdfplumber"):
                structure_data = self._extract_with_pdfplumber(file_path)
            else:
                structure_data = self._extract_with_pypdf2(file_path)

        except Exception as e:
            logger.error(f"PDF structure extraction failed: {e}")
            structure_data["metadata"]["extraction_error"] = str(e)

        return DocumentStructure(
            headers=structure_data["headers"],
            tables=structure_data["tables"],
            paragraphs=structure_data["paragraphs"],
            lists=structure_data["lists"],
            metadata=structure_data["metadata"],
            formatting=structure_data["formatting"],
            structure_type="pdf",
        )

    def _extract_with_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """Extract structure using PyPDF2."""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)

            full_text = ""
            page_texts = []

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                page_texts.append(page_text)
                full_text += page_text + "\n\n"

            # Extract text-based structure
            headers = self.text_analyzer.extract_headers(full_text)
            lists = self.text_analyzer.extract_lists(full_text)
            paragraphs = self.text_analyzer.extract_paragraphs(full_text)

            return {
                "headers": headers,
                "tables": [],  # Limited table extraction with PyPDF2
                "paragraphs": paragraphs,
                "lists": lists,
                "metadata": {
                    "page_count": len(reader.pages),
                    "extraction_method": "pypdf2",
                },
                "formatting": {"page_texts": page_texts},
            }


class DOCXStructureExtractor:
    """Extract structure from DOCX documents."""

    def __init__(self):
        self.text_analyzer = TextAnalyzer()

    def extract_structure(self, file_path: str) -> DocumentStructure:
        """Extract complete structure from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")

        try:
            doc = DocxDocument(file_path)

            headers = self._extract_docx_headers(doc)
            tables = self._extract_docx_tables(doc)
            paragraphs = self._extract_docx_paragraphs(doc)
            lists = self._extract_docx_lists(doc)

            return DocumentStructure(
                headers=headers,
                tables=tables,
                paragraphs=paragraphs,
                lists=lists,
                metadata={
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "extraction_method": "python-docx",
                },
                formatting=self._extract_docx_formatting(doc),
                structure_type="docx",
            )

        except Exception as e:
            logger.error(f"DOCX structure extraction failed: {e}")
            return DocumentStructure(
                headers=[],
                tables=[],
                paragraphs=[],
                lists=[],
                metadata={"extraction_error": str(e)},
                formatting={},
                structure_type="docx",
            )

    def _extract_docx_headers(self, doc: DocxDocument) -> List[HeaderInfo]:
        """Extract headers from DOCX using styles."""
        headers = []

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name.startswith("Heading"):
                try:
                    level = int(paragraph.style.name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                headers.append(HeaderInfo(level=level, text=paragraph.text, position=i))

        return self.text_analyzer._build_header_hierarchy(headers)

    def _extract_docx_tables(self, doc: DocxDocument) -> List[TableData]:
        """Extract tables from DOCX."""
        tables = []

        for table in doc.tables:
            rows = []
            headers = None

            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]

                if i == 0 and self._looks_like_header(row_data):
                    headers = row_data
                else:
                    rows.append(row_data)

            if rows or headers:
                tables.append(
                    TableData(rows=rows, headers=headers, table_type="docx_table")
                )

        return tables

    def _extract_docx_paragraphs(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract paragraphs from DOCX with formatting."""
        paragraphs = []

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() and not paragraph.style.name.startswith(
                "Heading"
            ):

                paragraphs.append(
                    {
                        "text": paragraph.text.strip(),
                        "position": i,
                        "style": paragraph.style.name,
                        "alignment": (
                            str(paragraph.alignment) if paragraph.alignment else None
                        ),
                    }
                )

        return paragraphs

    def _extract_docx_lists(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract lists from DOCX."""
        # This is a simplified version - DOCX list extraction is complex
        text = "\n".join([p.text for p in doc.paragraphs])
        return self.text_analyzer.extract_lists(text)

    def _extract_docx_formatting(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract formatting information from DOCX."""
        styles_used = set()
        for paragraph in doc.paragraphs:
            styles_used.add(paragraph.style.name)

        return {"styles_used": list(styles_used), "has_tables": len(doc.tables) > 0}

    def _looks_like_header(self, row_data: List[str]) -> bool:
        """Determine if a table row looks like a header."""
        if not row_data:
            return False

        # Check if all cells have content and are relatively short
        for cell in row_data:
            if not cell.strip() or len(cell.split()) > 5:
                return False

        return True


class HTMLStructureExtractor:
    """Extract structure from HTML documents."""

    def extract_structure(self, file_path: str) -> DocumentStructure:
        """Extract complete structure from HTML."""
        if not HTML_AVAILABLE:
            raise ImportError("BeautifulSoup4 not available")

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()

            soup = BeautifulSoup(content, "html.parser")

            headers = self._extract_html_headers(soup)
            tables = self._extract_html_tables(soup)
            paragraphs = self._extract_html_paragraphs(soup)
            lists = self._extract_html_lists(soup)

            return DocumentStructure(
                headers=headers,
                tables=tables,
                paragraphs=paragraphs,
                lists=lists,
                metadata={
                    "title": soup.title.string if soup.title else None,
                    "extraction_method": "beautifulsoup",
                },
                formatting=self._extract_html_formatting(soup),
                structure_type="html",
            )

        except Exception as e:
            logger.error(f"HTML structure extraction failed: {e}")
            return DocumentStructure(
                headers=[],
                tables=[],
                paragraphs=[],
                lists=[],
                metadata={"extraction_error": str(e)},
                formatting={},
                structure_type="html",
            )

    def _extract_html_headers(self, soup: BeautifulSoup) -> List[HeaderInfo]:
        """Extract headers from HTML."""
        headers = []
        all_header_tags = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])

        for i, tag in enumerate(all_header_tags):
            level = int(tag.name[1])
            headers.append(
                HeaderInfo(level=level, text=tag.get_text(strip=True), position=i)
            )

        return headers

    def _extract_html_tables(self, soup: BeautifulSoup) -> List[TableData]:
        """Extract tables from HTML."""
        tables = []

        for table in soup.find_all("table"):
            rows = []
            headers = None

            # Look for header row
            thead = table.find("thead")
            if thead:
                header_row = thead.find("tr")
                if header_row:
                    headers = [
                        th.get_text(strip=True)
                        for th in header_row.find_all(["th", "td"])
                    ]

            # Extract data rows
            tbody = table.find("tbody") or table
            for row in tbody.find_all("tr"):
                cells = row.find_all(["td", "th"])
                if cells:
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    if (
                        not headers
                        and row == tbody.find("tr")
                        and all(cell.name == "th" for cell in cells)
                    ):
                        headers = row_data
                    else:
                        rows.append(row_data)

            if rows or headers:
                # Look for caption
                caption = table.find("caption")
                caption_text = caption.get_text(strip=True) if caption else None

                tables.append(
                    TableData(
                        rows=rows,
                        headers=headers,
                        caption=caption_text,
                        table_type="html_table",
                    )
                )

        return tables

    def _extract_html_paragraphs(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """Extract paragraphs from HTML."""
        paragraphs = []

        for i, p in enumerate(soup.find_all("p")):
            text = p.get_text(strip=True)
            if text:
                paragraphs.append(
                    {
                        "text": text,
                        "position": i,
                        "classes": p.get("class", []),
                        "id": p.get("id"),
                    }
                )

        return paragraphs

    def _extract_html_lists(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """Extract lists from HTML."""
        lists = []

        for list_tag in soup.find_all(["ul", "ol"]):
            list_type = "numbered" if list_tag.name == "ol" else "bullet"
            items = [
                li.get_text(strip=True)
                for li in list_tag.find_all("li", recursive=False)
            ]

            if items:
                lists.append(
                    {
                        "type": list_type,
                        "items": items,
                        "classes": list_tag.get("class", []),
                        "id": list_tag.get("id"),
                    }
                )

        return lists

    def _extract_html_formatting(self, soup: BeautifulSoup) -> Dict[str, Any]:
        """Extract formatting information from HTML."""
        return {
            "has_css": bool(soup.find_all(["style", "link"])),
            "has_scripts": bool(soup.find_all("script")),
            "meta_tags": {
                meta.get("name"): meta.get("content")
                for meta in soup.find_all("meta")
                if meta.get("name")
            },
            "image_count": len(soup.find_all("img")),
            "link_count": len(soup.find_all("a")),
        }


# Main structure extraction function
def extract_document_structure(file_path: str) -> DocumentStructure:
    """Extract structure from any supported document format."""

    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf":
            extractor = PDFStructureExtractor()
            return extractor.extract_structure(file_path)
        elif ext == ".docx":
            extractor = DOCXStructureExtractor()
            return extractor.extract_structure(file_path)
        elif ext in [".html", ".htm"]:
            extractor = HTMLStructureExtractor()
            return extractor.extract_structure(file_path)
        else:
            # Fallback to text analysis
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()

            analyzer = TextAnalyzer()
            return DocumentStructure(
                headers=analyzer.extract_headers(text),
                tables=[],
                paragraphs=analyzer.extract_paragraphs(text),
                lists=analyzer.extract_lists(text),
                metadata={"extraction_method": "text_analysis"},
                formatting={},
                structure_type="text",
            )

    except Exception as e:
        logger.error(f"Structure extraction failed for {file_path}: {e}")
        return DocumentStructure(
            headers=[],
            tables=[],
            paragraphs=[],
            lists=[],
            metadata={"extraction_error": str(e)},
            formatting={},
            structure_type="error",
        )


if __name__ == "__main__":
    # Example usage
    test_file = "sample_document.pdf"
    if Path(test_file).exists():
        structure = extract_document_structure(test_file)
        print(f"Headers: {len(structure.headers)}")
        print(f"Tables: {len(structure.tables)}")
        print(f"Paragraphs: {len(structure.paragraphs)}")
        print(f"Lists: {len(structure.lists)}")
